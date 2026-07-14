#!/usr/bin/env python3
"""Convert BRD md_log to .docx (python-docx).

Usage:
    python3 brd-to-docx.py < md_log.md          # stdin
    python3 brd-to-docx.py path/to/md_log.md    # file path
    python3 brd-to-docx.py -o output.docx < md_log.md

Output: bt_<date>_<UUID>.docx (or specified -o)
"""

import sys
import re
import os
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


def parse_md_log(text: str) -> dict:
    """Parse BRD md_log into structured dict."""
    data = {
        "meta": {},
        "session": "",
        "question_log": [],
        "summary": {},
        "compiled": {},
        "verification": {}
    }

    lines = text.split("\n")
    section = None
    current_q = None
    q_lines = []

    for line in lines:
        # Session title
        m = re.match(r'^# BRD Session:\s*(\S+)', line)
        if m:
            data["session"] = m.group(1)
            continue

        # Section headers
        if line.startswith("## Meta"):
            section = "meta"
            continue
        if line.startswith("## QuestionLog"):
            section = "question_log"
            continue
        if line.startswith("## Summary"):
            section = "summary"
            continue
        if line.startswith("## CompiledBRD"):
            section = "compiled"
            continue
        if line.startswith("## VerificationLog"):
            section = "verification"
            continue

        # Meta key: value
        if section == "meta":
            m = re.match(r'^-\s*(\w+):\s*(.*)', line)
            if m:
                data["meta"][m.group(1)] = m.group(2).strip()

        # QuestionLog: Q1, Q2, etc.
        if section == "question_log":
            m = re.match(r'^###\s*Q(\d+):\s*(.*)', line)
            if m:
                if current_q:
                    data["question_log"].append(current_q)
                current_q = {"num": m.group(1), "question": m.group(2), "answer": ""}
                continue
            m = re.match(r'^\*\*A:\*\*\s*(.*)', line)
            if m and current_q:
                current_q["answer"] = m.group(1)
                continue

        # Summary
        if section == "summary":
            m = re.match(r'^-\s*\*\*?(Симптом|Барьер|Потери|Контекст)\*?\*?:\s*(.*)', line)
            if m:
                data["summary"][m.group(1)] = m.group(2).strip()
            m = re.match(r'^-\s*(Симптом|Барьер|Потери|Контекст):\s*(.*)', line)
            if m:
                data["summary"][m.group(1)] = m.group(2).strip()

        # CompiledBRD
        if section == "compiled":
            if line.startswith("### Description"):
                section = "compiled_desc"
                continue
            if line.startswith("### Goal"):
                section = "compiled_goal"
                continue
            if line.startswith("### Metrics"):
                section = "compiled_metrics"
                continue
            if line.startswith("### Impacts"):
                section = "compiled_impacts"
                continue
            if line.startswith("### References"):
                section = "compiled_refs"
                continue

        if section == "compiled_desc":
            if line.strip() and not line.startswith("###"):
                data["compiled"]["description"] = (data["compiled"].get("description", "") + " " + line.strip()).strip()
            elif line.startswith("###"):
                section = "compiled"

        if section == "compiled_goal":
            if line.strip() and not line.startswith("###"):
                data["compiled"]["goal"] = (data["compiled"].get("goal", "") + " " + line.strip()).strip()
            elif line.startswith("###"):
                section = "compiled"

        if section == "compiled_metrics":
            if line.strip() and not line.startswith("###"):
                data["compiled"]["metrics"] = (data["compiled"].get("metrics", "") + "\n" + line.strip()).strip()
            elif line.startswith("###"):
                section = "compiled"

        if section == "compiled_impacts":
            if line.strip() and not line.startswith("###"):
                data["compiled"]["impacts"] = (data["compiled"].get("impacts", "") + "\n" + line.strip()).strip()
            elif line.startswith("###"):
                section = "compiled"

        if section == "compiled_refs":
            m = re.match(r'^-\s*\[(.+)\]\((.+)\)', line)
            if m:
                if "refs" not in data["compiled"]:
                    data["compiled"]["refs"] = []
                data["compiled"]["refs"].append({"title": m.group(1), "url": m.group(2)})
            elif line.strip() and not line.startswith("###"):
                data["compiled"]["refs_extra"] = (data["compiled"].get("refs_extra", "") + "\n" + line.strip()).strip()

        # VerificationLog
        if section == "verification":
            if line.startswith("### Check:"):
                continue
            m = re.match(r'^- Status:\s*([✅❌🔄]+)', line)
            if m:
                status = "PASS" if "✅" in m.group(1) else ("FAIL" if "❌" in m.group(1) else "RETRY")
                data["verification"]["status"] = status
                continue
            m = re.match(r'^- Comment:\s*(.*)', line)
            if m:
                data["verification"]["comment"] = m.group(1).strip()
                continue
            m = re.match(r'^- Verdict:\s*(.*)', line)
            if m:
                data["verification"]["verdict"] = m.group(1).strip()

    # Last question
    if current_q:
        data["question_log"].append(current_q)

    # Raw text capture for sections we can't parse
    data["raw_compiled"] = ""
    in_compiled = False
    in_verification = False
    compiled_section = None
    for line in lines:
        if line.startswith("## CompiledBRD"):
            in_compiled = True
            in_verification = False
            continue
        if line.startswith("## VerificationLog"):
            in_verification = True
            in_compiled = False
            continue
        if line.startswith("## "):
            in_compiled = False
            in_verification = False
            continue
        if in_compiled:
            data["raw_compiled"] += line + "\n"
        if in_verification:
            m = re.match(r'^- Verdict:\s*(.*)', line)
            if m:
                data["verification"]["verdict"] = m.group(1).strip()
            m = re.match(r'### Check:\s*(.*)', line)
            if m:
                data["verification"][f"check_{m.group(1).strip()}"] = m.group(1).strip()
            m = re.match(r'^- Status:\s*([✅❌🔄]+)', line)
            if m:
                status_map = {"✅": "PASS", "❌": "FAIL", "🔄": "RETRY"}
                data["verification"]["rca_status"] = status_map.get(m.group(1), m.group(1))
            m = re.match(r'^- Comment:\s*(.*)', line)
            if m:
                data["verification"]["comment"] = m.group(1).strip()

    return data


def generate_docx(data: dict, output_path: str):
    """Generate .docx from parsed data."""
    doc = Document()

    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title = doc.add_heading('Business Requirements Document (BRD)', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Session info
    p = doc.add_paragraph()
    p.add_run(f'Session: ').bold = True
    p.add_run(data["session"])
    p = doc.add_paragraph()
    p.add_run(f'Date: ').bold = True
    p.add_run(datetime.now().strftime("%Y-%m-%d %H:%M"))
    if data["meta"].get("user"):
        p.add_run(f' | User: ')
        p.add_run(data["meta"]["user"]).bold = True

    doc.add_page_break()

    # 1. Problem Description
    doc.add_heading('1. Root Cause Analysis', level=1)
    raw = data.get("raw_compiled", "")
    # Extract section 1 (RCA) — everything until next numbered section or end
    rca_match = re.search(r'###\s*1\.\s*Описание проблемы.*?\n(.*?)(?=\n###\s*2\.|\Z)', raw, re.DOTALL)
    if rca_match:
        rca_text = rca_match.group(1).strip()
        rca_text = re.sub(r'\*\*', '', rca_text)
        for para in rca_text.split('\n'):
            if para.strip():
                doc.add_paragraph(para.strip())
    elif data["summary"]:
        for key, val in data["summary"].items():
            p = doc.add_paragraph()
            p.add_run(f'{key}: ').bold = True
            p.add_run(val)
    else:
        doc.add_paragraph("—")

    # 2. Goal
    doc.add_heading('2. SMART-цель', level=1)
    goal_match = re.search(r'###\s*2\.\s*SMART-цель.*?\n(.*?)(?=\n###\s*3\.|\Z)', raw, re.DOTALL)
    if goal_match:
        goal_text = goal_match.group(1).strip()
        goal_text = re.sub(r'\*\*', '', goal_text)
        for para in goal_text.split('\n'):
            if para.strip():
                doc.add_paragraph(para.strip())
    else:
        doc.add_paragraph("—")

    # 3. KPI
    doc.add_heading('3. Ключевые показатели (KPI)', level=1)
    kpi_match = re.search(r'###\s*3\.\s*Метрики успеха.*?\n(.*?)(?=\n###\s*4\.|\Z)', raw, re.DOTALL)
    if kpi_match:
        kpi_text = kpi_match.group(1).strip()
        # Try to detect a markdown table and render it
        table_lines = []
        in_table = False
        for line in kpi_text.split('\n'):
            if '|' in line:
                if not in_table:
                    in_table = True
                table_lines.append(line)
            else:
                if in_table:
                    # render table
                    if len(table_lines) >= 2:
                        # parse header
                        header = [h.strip() for h in table_lines[0].strip('|').split('|')]
                        # skip delimiter row
                        rows = []
                        for tl in table_lines[2:]:
                            cells = [c.strip() for c in tl.strip('|').split('|')]
                            if cells:
                                rows.append(cells)
                        if rows:
                            tbl = doc.add_table(rows=len(rows)+1, cols=len(header))
                            tbl.style = 'Light Grid Accent 1'
                            for i, h in enumerate(header):
                                tbl.cell(0, i).text = h
                            for ri, row in enumerate(rows):
                                for ci, cell in enumerate(row):
                                    if ci < len(header):
                                        tbl.cell(ri+1, ci).text = cell
                    table_lines = []
                    in_table = False
                if line.strip():
                    doc.add_paragraph(line.strip())
        # handle trailing table
        if in_table and len(table_lines) >= 2:
            header = [h.strip() for h in table_lines[0].strip('|').split('|')]
            rows = []
            for tl in table_lines[2:]:
                cells = [c.strip() for c in tl.strip('|').split('|')]
                if cells and len(cells) > 1:
                    rows.append(cells)
            if rows:
                tbl = doc.add_table(rows=len(rows)+1, cols=len(header))
                tbl.style = 'Light Grid Accent 1'
                for i, h in enumerate(header):
                    tbl.cell(0, i).text = h
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        if ci < len(header):
                            tbl.cell(ri+1, ci).text = cell
    else:
        doc.add_paragraph("—")

    # 4. Solution
    doc.add_heading('4. Решение: Мера.НПА', level=1)
    sol_match = re.search(r'###\s*4\.\s*Описание желаемого решения.*?\n(.*?)(?=\n###\s*5\.|\Z)', raw, re.DOTALL)
    if sol_match:
        sol_text = sol_match.group(1).strip()
        sol_text = re.sub(r'\*\*', '', sol_text)
        for para in sol_text.split('\n'):
            if para.strip():
                doc.add_paragraph(para.strip())
    else:
        doc.add_paragraph("—")

    # 5. Requirements
    doc.add_heading('5. Бизнес-требования', level=1)
    req_match = re.search(r'###\s*5\.\s*Бизнес-требования.*?\n(.*?)(?=\n###\s*6\.|\Z)', raw, re.DOTALL)
    if req_match:
        req_text = req_match.group(1).strip()
        req_text = re.sub(r'\*\*', '', req_text)
        for para in req_text.split('\n'):
            if para.strip():
                doc.add_paragraph(para.strip())
    # Also include interview questions (Q&A)
    doc.add_heading('5.1 Результаты интервью', level=2)
    if data["question_log"]:
        for q in data["question_log"]:
            p = doc.add_paragraph()
            p.add_run(f'Вопрос {q["num"]}: ').bold = True
            p.add_run(q["question"])
            p2 = doc.add_paragraph()
            p2.add_run('Ответ: ').bold = True
            p2.add_run(q.get("answer", "—"))
    else:
        doc.add_paragraph("—")

    # 6. Effect
    doc.add_heading('6. Оценка эффекта (до/после)', level=1)
    eff_match = re.search(r'###\s*6\.\s*Оценка эффекта.*?\n(.*?)(?=\n###\s*7\.|\Z|\n##\s)', raw, re.DOTALL)
    if eff_match:
        eff_text = eff_match.group(1).strip()
        # Parse effect table
        table_lines = []
        in_table = False
        for line in eff_text.split('\n'):
            if '|' in line and not line.strip().startswith('|--'):
                if not in_table:
                    in_table = True
                table_lines.append(line)
            elif '|' in line and line.strip().startswith('|--'):
                continue
            else:
                if in_table and len(table_lines) >= 2:
                    header = [h.strip() for h in table_lines[0].strip('|').split('|')]
                    rows = []
                    for tl in table_lines[1:]:
                        cells = [c.strip() for c in tl.strip('|').split('|')]
                        if cells and len(cells) > 1:
                            rows.append(cells)
                    if rows:
                        tbl = doc.add_table(rows=len(rows)+1, cols=len(header))
                        tbl.style = 'Light Grid Accent 1'
                        for i, h in enumerate(header):
                            tbl.cell(0, i).text = h
                        for ri, row in enumerate(rows):
                            for ci, cell in enumerate(row):
                                if ci < len(header):
                                    tbl.cell(ri+1, ci).text = cell
                    table_lines = []
                    in_table = False
                if line.strip():
                    p = doc.add_paragraph()
                    # strip **bold** markers
                    p.add_run(line.strip().replace('**', ''))
        # handle trailing table
        if in_table and len(table_lines) >= 2:
            header = [h.strip() for h in table_lines[0].strip('|').split('|')]
            rows = []
            for tl in table_lines[1:]:
                cells = [c.strip() for c in tl.strip('|').split('|')]
                if cells and len(cells) > 1:
                    rows.append(cells)
            if rows:
                tbl = doc.add_table(rows=len(rows)+1, cols=len(header))
                tbl.style = 'Light Grid Accent 1'
                for i, h in enumerate(header):
                    tbl.cell(0, i).text = h
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        if ci < len(header):
                            tbl.cell(ri+1, ci).text = cell
    else:
        doc.add_paragraph("—")

    # 7. Verification
    doc.add_heading('7. Верификация', level=1)
    v = data["verification"]
    if v.get("verdict"):
        p = doc.add_paragraph()
        p.add_run('Вердикт: ').bold = True
        p.add_run(v["verdict"])
    if v.get("rca_status"):
        p = doc.add_paragraph()
        p.add_run(f'RCA Integrity: ')
        p.add_run(v["rca_status"])
    if v.get("comment"):
        p = doc.add_paragraph()
        p.add_run('Комментарий: ').bold = True
        p.add_run(v["comment"])

    if not v:
        doc.add_paragraph("Верификация не пройдена.")

    doc.save(output_path)
    return output_path


def main():
    output_path = None

    # Parse args
    args = sys.argv[1:]
    input_source = None  # file path or "-" or None (stdin)

    for i, arg in enumerate(args):
        if arg == "-o" and i + 1 < len(args):
            output_path = args[i + 1]
        elif arg.startswith("-"):
            continue
        elif input_source is None:
            input_source = arg

    # Read input
    if input_source and input_source != "-":
        with open(input_source, "r", encoding="utf-8") as f:
            text = f.read()
    else:
        text = sys.stdin.read()

    if not text.strip():
        print("ERROR: Empty input")
        sys.exit(1)

    # Parse
    data = parse_md_log(text)

    # Generate filename if not specified
    if not output_path:
        date_str = datetime.now().strftime("%Y%m%d")
        session_id = data.get("session", "unknown")
        output_path = f"bt_{date_str}_{session_id}.docx"

    # Generate docx
    result = generate_docx(data, output_path)
    print(f"✅ BRD saved: {result}")


if __name__ == "__main__":
    main()
