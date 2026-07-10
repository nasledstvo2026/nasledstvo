#!/usr/bin/env python3
"""
gen.py — Генератор бизнес-требований (БТ)
Версия 2.0, 09.07.2026

Принимает JSON с ответами на 7 блоков → генерирует DOCX + HTML,
публикует на GitHub Pages.

Использование:
  python3 gen.py --answers answers.json    # из файла
  python3 gen.py --help                     # справка

Формат answers.json:
{
  "title": "Название инициативы",
  "sections": [
    { "id": 1, "name": "Контекст и проблема",
      "answers": {
        "1.1": "Рабочее название: Калькулятор наследства",
        "1.2": "Пользователи не знают...",
        "1.3": "Теряем по 100 обращений в месяц..."
      }
    },
    ...
  ]
}
"""

import json
import sys
import os
import argparse
import subprocess
import tempfile
from datetime import datetime

# ------------------------------
# Шаблон БТ (сокращённая версия)
# ------------------------------
TEMPLATE_PATH = os.path.join(os.path.dirname(__file__),
                             '_bmad-output/planning-artifacts/template-7-blocks.json')


def load_template():
    """Загружает шаблон 7 блоков из template-7-blocks.json"""
    with open(TEMPLATE_PATH, encoding='utf-8') as f:
        return json.load(f)


def build_docx_text(answers, template):
    """Собирает текст DOCX из ответов"""
    lines = []
    lines.append(f'# {answers.get("title", "Бизнес-требования")}')
    lines.append(f'Дата: {datetime.now().strftime("%d.%m.%Y")}')
    lines.append('')

    for section in answers.get('sections', []):
        block_id = str(section['id'])
        block_name = section['name']
        lines.append(f'## {block_id}. {block_name}')
        lines.append('')

        for subsection in section.get('answers', {}):
            answer = section['answers'][subsection]
            # Ищем название подблока в шаблоне
            sub_name = _find_subsection_name(template, block_id, subsection)
            if sub_name:
                lines.append(f'### {subsection}. {sub_name}')
            lines.append(answer)
            lines.append('')

    return '\n'.join(lines)


def _find_subsection_name(template, block_id, sub_id):
    """Ищет название подраздела по id"""
    for block in template.get('blocks', []):
        if str(block['id']) == block_id:
            for sub in block.get('subsections', []):
                if sub['id'] == sub_id:
                    return sub['name']
    return sub_id


def generate_html_text(docx_text, title):
    """Генерирует HTML-страницу из текста DOCX"""
    lines = []
    lines.append('<!DOCTYPE html>')
    lines.append('<html lang="ru">')
    lines.append('<head>')
    lines.append('<meta charset="UTF-8">')
    lines.append('<meta name="viewport" content="width=device-width, initial-scale=1.0">')
    lines.append(f'<title>{title} — Бизнес-требования</title>')
    lines.append('<link rel="stylesheet" href="theme.css">')
    lines.append('</head>')
    lines.append('<body>')
    lines.append('<div class="container report">')
    lines.append(f'<div class="hero"><h1>{title}</h1></div>')

    in_list = False
    for line in docx_text.split('\n'):
        if line.startswith('# '):
            continue  # заголовок уже в hero
        elif line.startswith('## '):
            if in_list:
                lines.append('</ul>')
                in_list = False
            lines.append(f'<h2>{line[3:].strip()}</h2>')
        elif line.startswith('### '):
            if in_list:
                lines.append('</ul>')
                in_list = False
            lines.append(f'<h3>{line[4:].strip()}</h3>')
        elif line.strip():
            if not in_list:
                lines.append('<ul>')
                in_list = True
            lines.append(f'<li>{line.strip()}</li>')
        else:
            if in_list:
                lines.append('</ul>')
                in_list = False

    if in_list:
        lines.append('</ul>')

    lines.append('<div class="footer">')
    lines.append(f'<p>Сгенерировано: {datetime.now().strftime("%d.%m.%Y %H:%M")}</p>')
    lines.append('</div>')
    lines.append('</div></body></html>')
    return '\n'.join(lines)


def generate_docx(docx_text, output_path):
    """Генерирует DOCX-файл"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
    except ImportError:
        print("❌ python-docx не установлен. Установи: pip install python-docx")
        sys.exit(1)

    doc = Document()

    # Стили
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    for line in docx_text.split('\n'):
        if line.startswith('## '):
            p = doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('# '):
            p = doc.add_heading(line[2:].strip(), level=1)
        elif line.strip():
            doc.add_paragraph(line.strip())
        # Пустые строки пропускаем

    doc.save(output_path)
    return output_path


def publish_to_github(file_path, filename):
    """Публикует файл через upload-to-github.sh"""
    script = os.path.join(os.path.dirname(__file__), 'upload-to-github.sh')
    if not os.path.exists(script):
        print(f"⚠️ upload-to-github.sh не найден ({script})")
        print(f"📄 Файл сохранён: {file_path}")
        return f"file://{file_path}"

    result = subprocess.run(
        ['bash', script, file_path, filename],
        capture_output=True, text=True
    )
    if result.returncode != 0:
        print(f"❌ Ошибка публикации: {result.stderr}")
        return f"file://{file_path}"

    url = result.stdout.strip()
    return url


def main():
    parser = argparse.ArgumentParser(description='Генератор бизнес-требований')
    parser.add_argument('--answers', required=True, help='JSON-файл с ответами')
    parser.add_argument('--output-dir', default='/tmp/bt-output',
                        help='Директория для выходных файлов')
    parser.add_argument('--no-publish', action='store_true',
                        help='Не публиковать на GitHub Pages')
    args = parser.parse_args()

    # Загружаем ответы
    with open(args.answers, encoding='utf-8') as f:
        answers = json.load(f)

    title = answers.get('title', 'Бизнес-требования')
    safe_title = title.replace(' ', '_').replace('/', '_')

    template = load_template()

    # Генерируем текст
    docx_text = build_docx_text(answers, template)
    html_text = generate_html_text(docx_text, title)

    # Создаём выходную директорию
    os.makedirs(args.output_dir, exist_ok=True)

    # Сохраняем DOCX
    docx_path = os.path.join(args.output_dir, f'{safe_title}.docx')
    generate_docx(docx_text, docx_path)
    print(f'✅ DOCX: {docx_path}')

    # Сохраняем HTML
    html_path = os.path.join(args.output_dir, f'{safe_title}.html')
    with open(html_path, 'w', encoding='utf-8') as f:
        f.write(html_text)
    print(f'✅ HTML: {html_path}')

    # Публикация
    if not args.no_publish:
        print('📤 Публикую на GitHub Pages...')
        docx_url = publish_to_github(docx_path, f'{safe_title}.docx')
        html_url = publish_to_github(html_path, f'{safe_title}.html')
        if docx_url:
            print(f'🌐 DOCX: {docx_url}')
        if html_url:
            print(f'🌐 HTML: {html_url}')

    print('\n✅ Готово!')


if __name__ == '__main__':
    main()
