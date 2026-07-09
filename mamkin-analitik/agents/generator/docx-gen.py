#!/usr/bin/env python3
"""
DOCX Generator — генерация DOCX-файла бизнес-требований.

Story 4.1 — Generator: создание профессионального DOCX-документа
с титульной страницей, оглавлением, 7 блоками и секцией рисков.

Использует python-docx (>=1.1.2) для создания профессионального документа
с оглавлением, стилями и форматированием.

Зависимости: pip install python-docx

API:
    generate_brd(data, output_path) — основная функция
    main() — CLI-точка входа

Формат входных данных (JSON):
    {
        "title": "Документ бизнес-требований",
        "createdAt": "2026-07-08T20:00:00.000Z",
        "sessionId": "session_xxx",
        "telegramUserId": 12345,
        "totalBlocks": 7,
        "blocks": [
            {
                "blockId": 1,
                "blockName": "Предпосылки и цели",
                "content": "## Блок 1. ...",
                "subsections": [
                    {
                        "subsectionId": "1.1",
                        "subsectionName": "Название",
                        "text": "— Ответ...",
                        "depthReached": "L2"
                    }
                ]
            }
        ],
        "risks": {
            "total": 5,
            "byCategory": { "technical": [...], "org": [...] },
            "items": [
                {
                    "text": "Описание риска",
                    "category": "technical",
                    "probability": 0.7,
                    "impact": 0.6,
                    "mitigation": "Митигация"
                }
            ]
        },
        "completedSubsections": 20,
        "totalSubsections": 21,
        "fullText": "Полный текст документа..."
    }
"""

import json
import sys
import os
import logging
import traceback
from datetime import datetime
from typing import Optional, Dict, Any, List

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ==================== Logging ====================

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] docx-gen: %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger('docx-gen')


# ==================== Constants ====================

FONT_FAMILY = 'Calibri'
FONT_FAMILY_HEADING = 'Calibri Light'
FONT_SIZE_BODY = Pt(11)
FONT_SIZE_TITLE = Pt(26)
FONT_SIZE_SUBTITLE = Pt(14)
FONT_SIZE_H1 = Pt(18)
FONT_SIZE_H2 = Pt(15)
FONT_SIZE_H3 = Pt(13)
FONT_SIZE_TOC = Pt(11)
FONT_SIZE_SMALL = Pt(9)

COLOR_PRIMARY = RGBColor(0x1F, 0x4E, 0x79)   # Dark blue
COLOR_ACCENT = RGBColor(0x2E, 0x75, 0xB6)     # Medium blue
COLOR_TEXT = RGBColor(0x33, 0x33, 0x33)        # Dark gray
COLOR_LIGHT = RGBColor(0x66, 0x66, 0x66)       # Gray
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_RISK_HIGH = RGBColor(0xC0, 0x39, 0x2B)   # Red
COLOR_RISK_MEDIUM = RGBColor(0xE6, 0x7E, 0x22)  # Orange
COLOR_RISK_LOW = RGBColor(0x27, 0xAE, 0x60)     # Green
COLOR_RISK_NONE = RGBColor(0x7F, 0x8C, 0x8D)    # Gray

MARGIN_TOP = Cm(2.0)
MARGIN_BOTTOM = Cm(2.0)
MARGIN_LEFT = Cm(2.5)
MARGIN_RIGHT = Cm(2.0)

MAX_RETRIES = 3


# ==================== Helper Functions ====================

def _set_cell_shading(cell, color_hex: str):
    """Устанавливает цвет фона ячейки таблицы."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _add_page_number(doc):
    """Добавляет номер страницы в нижний колонтитул."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run()
        run.font.size = FONT_SIZE_SMALL
        run.font.color.rgb = COLOR_LIGHT
        run.font.name = FONT_FAMILY

        # Add page number field
        fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fld_char_begin)

        run2 = p.add_run()
        instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._r.append(instr)

        run3 = p.add_run()
        fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._r.append(fld_char_end)

        # Add " / " and total pages
        run4 = p.add_run(' / ')
        run4.font.size = FONT_SIZE_SMALL
        run4.font.color.rgb = COLOR_LIGHT
        run4.font.name = FONT_FAMILY

        run5 = p.add_run()
        instr2 = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> NUMPAGES </w:instrText>')
        run5._r.append(instr2)

        run6 = p.add_run()
        fld_char_end2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run6._r.append(fld_char_end2)


def _setup_styles(doc):
    """Настраивает стили документа."""
    style = doc.styles['Normal']
    style.font.name = FONT_FAMILY
    style.font.size = FONT_SIZE_BODY
    style.font.color.rgb = COLOR_TEXT
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Heading 1
    if 'Heading 1' in doc.styles:
        h1 = doc.styles['Heading 1']
        h1.font.name = FONT_FAMILY_HEADING
        h1.font.size = FONT_SIZE_H1
        h1.font.color.rgb = COLOR_PRIMARY
        h1.font.bold = True
        h1.paragraph_format.space_before = Pt(24)
        h1.paragraph_format.space_after = Pt(12)
        h1.paragraph_format.keep_with_next = True

    # Heading 2
    if 'Heading 2' in doc.styles:
        h2 = doc.styles['Heading 2']
        h2.font.name = FONT_FAMILY_HEADING
        h2.font.size = FONT_SIZE_H2
        h2.font.color.rgb = COLOR_ACCENT
        h2.font.bold = True
        h2.paragraph_format.space_before = Pt(18)
        h2.paragraph_format.space_after = Pt(8)
        h2.paragraph_format.keep_with_next = True

    # Heading 3
    if 'Heading 3' in doc.styles:
        h3 = doc.styles['Heading 3']
        h3.font.name = FONT_FAMILY_HEADING
        h3.font.size = FONT_SIZE_H3
        h3.font.color.rgb = COLOR_TEXT
        h3.font.bold = True
        h3.paragraph_format.space_before = Pt(12)
        h3.paragraph_format.space_after = Pt(6)
        h3.paragraph_format.keep_with_next = True


def _add_paragraph(doc, text: str, style: str = None,
                   bold: bool = False, italic: bool = False,
                   size: Pt = None, color: RGBColor = None,
                   alignment: int = None, space_before: Pt = None,
                   space_after: Pt = None, font_name: str = None):
    """Добавляет параграф с форматированием."""
    p = doc.add_paragraph()
    if style:
        p.style = doc.styles[style]
    run = p.add_run(text)

    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = size
    if color:
        run.font.color.rgb = color
    if font_name:
        run.font.name = font_name
    if alignment is not None:
        p.alignment = alignment
    if space_before is not None:
        p.paragraph_format.space_before = space_before
    if space_after is not None:
        p.paragraph_format.space_after = space_after

    return p


def _add_rich_paragraph(doc, runs_config: List[Dict], alignment=None,
                        space_before=None, space_after=None):
    """Добавляет параграф с разными форматированиями внутри."""
    p = doc.add_paragraph()
    for cfg in runs_config:
        run = p.add_run(cfg.get('text', ''))
        if cfg.get('bold'):
            run.bold = True
        if cfg.get('italic'):
            run.italic = True
        if cfg.get('size'):
            run.font.size = cfg['size']
        if cfg.get('color'):
            run.font.color.rgb = cfg['color']
        if cfg.get('font_name'):
            run.font.name = cfg['font_name']

    if alignment is not None:
        p.alignment = alignment
    if space_before is not None:
        p.paragraph_format.space_before = space_before
    if space_after is not None:
        p.paragraph_format.space_after = space_after

    return p


def _format_risk_level(level: float) -> str:
    """Форматирует числовой уровень риска в текстовый."""
    if level is None:
        return 'Не определена'
    if level >= 0.7:
        return 'Высокая'
    if level >= 0.4:
        return 'Средняя'
    if level >= 0.1:
        return 'Низкая'
    return 'Не определена'


def _get_risk_color(level: float) -> RGBColor:
    """Возвращает цвет для уровня риска."""
    if level is None:
        return COLOR_RISK_NONE
    if level >= 0.7:
        return COLOR_RISK_HIGH
    if level >= 0.4:
        return COLOR_RISK_MEDIUM
    if level >= 0.1:
        return COLOR_RISK_LOW
    return COLOR_RISK_NONE


def _get_category_label(category: str) -> str:
    """Возвращает русскоязычную метку категории."""
    labels = {
        'technical': 'Технические риски',
        'org': 'Организационные риски',
        'organizational': 'Организационные риски',
        'business': 'Бизнес-риски',
        'adoption': 'Риски внедрения',
        'uncategorized': 'Прочие риски',
    }
    return labels.get(category, category)


def _extract_blocks_from_content(content: str) -> List[str]:
    """Извлекает текст блоков из markdown-контента."""
    blocks_text = []
    current_lines = []
    for line in content.split('\n'):
        if line.startswith('## Блок '):
            if current_lines:
                blocks_text.append('\n'.join(current_lines))
            current_lines = [line]
        else:
            current_lines.append(line)
    if current_lines:
        blocks_text.append('\n'.join(current_lines))
    return blocks_text


# ==================== Title Page ====================

def _build_title_page(doc, data: Dict):
    """Строит титульную страницу документа."""
    # Пустой параграф для отступа сверху
    for _ in range(6):
        _add_paragraph(doc, '', space_after=Pt(0))

    # Горизонтальная линия (через таблицу)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.cell(0, 0)
    cell.text = ''
    _set_cell_shading(cell, '1F4E79')
    cell.width = Inches(6.0)

    doc.add_paragraph()  # spacing

    # Название документа
    title = data.get('title', 'Документ бизнес-требований')
    _add_paragraph(
        doc, title,
        bold=True, size=FONT_SIZE_TITLE, color=COLOR_PRIMARY,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(8)
    )

    # Подзаголовок
    _add_paragraph(
        doc, 'Business Requirements Document',
        italic=True, size=FONT_SIZE_SUBTITLE, color=COLOR_LIGHT,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(30)
    )

    # Ещё одна линия
    table2 = doc.add_table(rows=1, cols=1)
    table2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell2 = table2.cell(0, 0)
    cell2.text = ''
    _set_cell_shading(cell2, '2E75B6')
    cell2.width = Inches(4.0)

    doc.add_paragraph()  # spacing

    # Метаданные
    meta_data = []

    # Дата создания
    created_at = data.get('createdAt', '')
    if created_at:
        try:
            dt = datetime.fromisoformat(created_at.replace('Z', '+00:00'))
            date_str = dt.strftime('%d.%m.%Y %H:%M')
        except (ValueError, AttributeError):
            date_str = created_at
        meta_data.append(('Дата создания', date_str))

    # Номер сессии
    meta_data.append(('Сессия', data.get('sessionId', '—')))

    # Автор / пользователь
    telegram_user = data.get('telegramUserId', '—')
    meta_data.append(('Автор (Telegram ID)', str(telegram_user)))

    # Количество блоков
    meta_data.append(('Количество разделов', str(data.get('totalBlocks', 7))))

    # Прогресс
    completed = data.get('completedSubsections', 0)
    total = data.get('totalSubsections', 0)
    meta_data.append(('Заполнено подразделов', f'{completed}/{total}'))

    # Риски
    risks = data.get('risks', {})
    meta_data.append(('Идентифицировано рисков', str(risks.get('total', 0))))

    # Выводим метаданные
    for label, value in meta_data:
        _add_rich_paragraph(
            doc,
            [
                {'text': label, 'bold': True, 'size': Pt(11), 'color': COLOR_PRIMARY},
                {'text': ':  ', 'size': Pt(11)},
                {'text': value, 'size': Pt(11), 'color': COLOR_TEXT},
            ],
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=Pt(4)
        )

    # Разрыв страницы
    doc.add_page_break()


# ==================== Table of Contents ====================

def _build_toc_page(doc, data: Dict):
    """Строит страницу оглавления."""
    _add_paragraph(
        doc, 'Содержание',
        bold=True, size=FONT_SIZE_H1, color=COLOR_PRIMARY,
        space_after=Pt(18)
    )

    # Статическое оглавление на основе блоков
    blocks = data.get('blocks', [])
    block_names = {
        1: 'Предпосылки и цели',
        2: 'Заинтересованные стороны',
        3: 'Текущее состояние',
        4: 'Требования',
        5: 'Критерии приёмки',
        6: 'Технические заметки',
        7: 'Риски'
    }

    for block in blocks:
        bid = block.get('blockId', 0)
        bname = block.get('blockName', block_names.get(bid, f'Блок {bid}'))
        subsections = block.get('subsections', [])

        # Заголовок блока
        _add_rich_paragraph(
            doc,
            [
                {'text': f'Блок {bid}. ', 'bold': True, 'size': Pt(12), 'color': COLOR_PRIMARY},
                {'text': bname, 'bold': True, 'size': Pt(12), 'color': COLOR_TEXT},
            ],
            space_before=Pt(10),
            space_after=Pt(4)
        )

        # Подразделы
        for sub in subsections:
            sub_id = sub.get('subsectionId', '')
            sub_name = sub.get('subsectionName', '')
            depth = sub.get('depthReached', '')

            depth_icon = {'L3': '●', 'L2': '◉', 'L1': '○', 'none': '○'}
            icon = depth_icon.get(depth, '○')
            depth_label = {'L3': ' (глубоко)', 'L2': ' (детально)', 'L1': ' (поверхностно)', 'none': ''}
            label = depth_label.get(depth, '')

            _add_rich_paragraph(
                doc,
                [
                    {'text': f'    {icon}  {sub_id}. {sub_name}', 'size': Pt(11), 'color': COLOR_TEXT},
                    {'text': label, 'size': Pt(9), 'color': COLOR_LIGHT, 'italic': True},
                ],
                space_after=Pt(2)
            )

    # Секция рисков (всегда в конце)
    risks = data.get('risks', {})
    risk_count = risks.get('total', 0)
    _add_rich_paragraph(
        doc,
        [
            {'text': 'Блок 7. Риски и митигации', 'bold': True, 'size': Pt(12), 'color': COLOR_PRIMARY},
            {'text': f'  ({risk_count} риск{"а" if risk_count % 10 in [2,3,4] and risk_count not in [12,13,14] else "ов" if risk_count != 1 else ""})', 'size': Pt(11), 'color': COLOR_TEXT},
        ],
        space_before=Pt(10),
        space_after=Pt(4)
    )

    doc.add_page_break()


# ==================== Blocks ====================

def _build_block_section(doc, block: Dict, block_number: int):
    """Строит один блок документа."""
    block_id = block.get('blockId', block_number)
    block_name = block.get('blockName', f'Блок {block_id}')
    subsections = block.get('subsections', [])

    # Заголовок блока (Heading 1)
    _add_paragraph(
        doc, f'Блок {block_id}. {block_name}',
        style='Heading 1',
        space_before=Pt(18),
        space_after=Pt(12)
    )

    # Подразделы
    for sub in subsections:
        sub_id = sub.get('subsectionId', '')
        sub_name = sub.get('subsectionName', '')
        text = sub.get('text', '')
        depth = sub.get('depthReached', '')

        # Заголовок подраздела (Heading 2)
        _add_paragraph(
            doc, f'{sub_id}. {sub_name}',
            style='Heading 2',
            space_before=Pt(14),
            space_after=Pt(6)
        )

        # Индикатор глубины
        depth_labels = {
            'L3': 'Глубокая проработка (L3)',
            'L2': 'Детальная проработка (L2)',
            'L1': 'Поверхностная проработка (L1)',
            'none': 'Не заполнено',
        }
        depth_label = depth_labels.get(depth, depth)
        depth_colors = {
            'L3': COLOR_RISK_LOW,
            'L2': COLOR_ACCENT,
            'L1': COLOR_RISK_MEDIUM,
            'none': COLOR_LIGHT,
        }
        depth_color = depth_colors.get(depth, COLOR_LIGHT)

        _add_paragraph(
            doc, f'Уровень проработки: {depth_label}',
            italic=True, size=FONT_SIZE_SMALL, color=depth_color,
            space_after=Pt(6)
        )

        # Текст подраздела
        if text:
            _render_formatted_text(doc, text)
        else:
            _add_paragraph(
                doc, '— Нет данных.',
                italic=True, color=COLOR_LIGHT,
                space_after=Pt(6)
            )


def _render_formatted_text(doc, text: str):
    """Рендерит форматированный текст (из compiler-а) в параграфы DOCX.

    Обрабатывает:
    - Маркированные списки (начинающиеся с —, -, *, •)
    - Нумерованные списки (1., 2. и т.д.)
    - Переносы строк внутри параграфа
    - **жирный текст**
    - *курсив*
    """
    if not text or text.strip() == '':
        return

    lines = text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        # Маркированный список
        if line.startswith('—') or line.startswith('- ') or line.startswith('* ') or line.startswith('•'):
            # Собираем все строки списка
            list_items = []
            while i < len(lines):
                l = lines[i].strip()
                if l.startswith('—') or l.startswith('- ') or l.startswith('* ') or l.startswith('•'):
                    # Удаляем маркер
                    item_text = l
                    for marker in ['— ', '- ', '* ', '• ']:
                        if item_text.startswith(marker):
                            item_text = item_text[len(marker):]
                            break
                    list_items.append(item_text)
                    i += 1
                elif l == '':
                    i += 1
                    break
                else:
                    break

            for item in list_items:
                _add_rich_paragraph(
                    doc,
                    [
                        {'text': '•  ', 'bold': True, 'color': COLOR_ACCENT, 'size': FONT_SIZE_BODY},
                        {'text': item, 'size': FONT_SIZE_BODY, 'color': COLOR_TEXT},
                    ],
                    space_after=Pt(3),
                    space_before=Pt(2)
                )
            continue

        # Нумерованный список
        if (len(line) > 2 and line[0].isdigit() and line[1] == '.'):
            list_items = []
            while i < len(lines):
                l = lines[i].strip()
                if len(l) > 2 and l[0].isdigit() and l[1] == '.':
                    # Извлекаем номер и текст
                    num_end = l.find('.')
                    num = l[:num_end]
                    item_text = l[num_end+1:].strip()
                    list_items.append((num, item_text))
                    i += 1
                elif l == '':
                    i += 1
                    break
                else:
                    break

            for num, item in list_items:
                _add_rich_paragraph(
                    doc,
                    [
                        {'text': f'{num}.  ', 'bold': True, 'color': COLOR_ACCENT, 'size': FONT_SIZE_BODY},
                        {'text': item, 'size': FONT_SIZE_BODY, 'color': COLOR_TEXT},
                    ],
                    space_after=Pt(3),
                    space_before=Pt(2)
                )
            continue

        # Обычный параграф
        _add_paragraph(
            doc, line,
            size=FONT_SIZE_BODY, color=COLOR_TEXT,
            space_after=Pt(6)
        )
        i += 1


# ==================== Risk Section ====================

def _build_risk_section(doc, risks: Dict):
    """Строит секцию рисков."""
    items = risks.get('items', [])
    by_category = risks.get('byCategory', {})

    if not items:
        _add_paragraph(
            doc, 'Риски не идентифицированы.',
            italic=True, color=COLOR_LIGHT,
            space_after=Pt(6)
        )
        return

    # Сводка рисков
    _add_paragraph(
        doc,
        f'Всего идентифицировано: {len(items)} риск(ов)',
        italic=True, size=FONT_SIZE_SMALL, color=COLOR_LIGHT,
        space_after=Pt(12)
    )

    # Таблица сводки рисков
    if items:
        risk_table = doc.add_table(rows=1, cols=5)
        risk_table.style = 'Table Grid'

        # Заголовки таблицы
        header_cells = risk_table.rows[0].cells
        headers = ['№', 'Риск', 'Категория', 'Вероятность', 'Влияние']
        for idx, header in enumerate(headers):
            header_cells[idx].text = header
            for paragraph in header_cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = FONT_SIZE_SMALL
                    run.font.color.rgb = COLOR_WHITE
                    run.font.name = FONT_FAMILY
            _set_cell_shading(header_cells[idx], '1F4E79')

        # Данные
        for idx, item in enumerate(items):
            row = risk_table.add_row()
            cells = row.cells

            risk_num = idx + 1
            risk_text = item.get('text', '—')
            if len(risk_text) > 100:
                risk_text = risk_text[:97] + '...'
            category = _get_category_label(item.get('category', 'uncategorized'))
            probability = item.get('probability')
            impact = item.get('impact')

            prob_text = _format_risk_level(probability)
            impact_text = _format_risk_level(impact)
            prob_color = _get_risk_color(probability)
            impact_color = _get_risk_color(impact)

            cells[0].text = str(risk_num)
            cells[1].text = risk_text
            cells[2].text = category
            cells[3].text = prob_text
            cells[4].text = impact_text

            for ci, cell in enumerate(cells):
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = FONT_SIZE_SMALL
                        run.font.name = FONT_FAMILY
                        if ci in (3, 4):
                            run.font.color.rgb = prob_color if ci == 3 else impact_color

            # Чередование фона строк
            if idx % 2 == 0:
                for cell in cells:
                    _set_cell_shading(cell, 'F2F7FB')

        doc.add_paragraph()  # spacing

    # Детальные риски по категориям
    for category, cat_items in by_category.items():
        cat_label = _get_category_label(category)
        _add_paragraph(
            doc, cat_label,
            style='Heading 2',
            space_before=Pt(16),
            space_after=Pt(8)
        )

        for idx, item in enumerate(cat_items):
            risk_text = item.get('text', '—')
            probability = item.get('probability')
            impact = item.get('impact')
            mitigation = item.get('mitigation')

            # Заголовок риска
            _add_rich_paragraph(
                doc,
                [
                    {'text': f'Риск {idx + 1}. ', 'bold': True, 'size': FONT_SIZE_BODY, 'color': COLOR_PRIMARY},
                    {'text': risk_text, 'size': FONT_SIZE_BODY, 'color': COLOR_TEXT},
                ],
                space_before=Pt(8),
                space_after=Pt(4)
            )

            # Детали риска (в таблице для компактности)
            details = []
            if probability is not None:
                prob_text = _format_risk_level(probability)
                details.append(('Вероятность', prob_text))
            if impact is not None:
                impact_text = _format_risk_level(impact)
                details.append(('Влияние', impact_text))

            if details:
                for label, value in details:
                    _add_rich_paragraph(
                        doc,
                        [
                            {'text': f'  {label}: ', 'bold': True, 'size': Pt(10), 'color': COLOR_LIGHT},
                            {'text': value, 'size': Pt(10), 'color': COLOR_TEXT},
                        ],
                        space_after=Pt(2)
                    )

            if mitigation:
                _add_rich_paragraph(
                    doc,
                    [
                        {'text': '  Митигация: ', 'bold': True, 'size': Pt(10), 'color': COLOR_RISK_LOW},
                        {'text': mitigation, 'size': Pt(10), 'color': COLOR_TEXT},
                    ],
                    space_after=Pt(6)
                )


# ==================== Main Generation Logic ====================

def generate_brd(data: dict, output_path: str) -> str:
    """
    Генерирует DOCX-документ БТ на основе данных сессии.

    Args:
        data: Словарь с данными БТ (блоки, ответы, риски)
        output_path: Путь для сохранения DOCX

    Returns:
        Путь к сохранённому файлу
    """
    logger.info(f"Starting DOCX generation → {output_path}")

    doc = Document()

    # Configure page margins
    for section in doc.sections:
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin = MARGIN_LEFT
        section.right_margin = MARGIN_RIGHT

    # Setup styles
    _setup_styles(doc)

    # 1. Title page
    logger.info("Building title page...")
    _build_title_page(doc, data)

    # 2. Table of contents
    logger.info("Building table of contents...")
    _build_toc_page(doc, data)

    # 3. Document body — 7 blocks
    blocks = data.get('blocks', [])
    logger.info(f"Building {len(blocks)} block(s)...")
    for block_number, block in enumerate(blocks, 1):
        _build_block_section(doc, block, block_number)

    # 4. Risk section
    risks = data.get('risks', {})
    doc.add_page_break()
    logger.info("Building risk section...")
    _add_paragraph(
        doc, 'Риски и митигации',
        style='Heading 1',
        space_before=Pt(6),
        space_after=Pt(12)
    )
    _build_risk_section(doc, risks)

    # 5. Page numbers
    _add_page_number(doc)

    # Save
    os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
    doc.save(output_path)
    logger.info(f"DOCX saved → {output_path}")

    return output_path


def generate_brd_with_retry(data: dict, output_path: str, max_retries: int = MAX_RETRIES) -> str:
    """
    Генерирует DOCX с повторными попытками при сбоях.

    Args:
        data: Данные сессии
        output_path: Путь сохранения
        max_retries: Максимальное количество попыток

    Returns:
        Путь к файлу
    """
    last_error = None

    for attempt in range(1, max_retries + 1):
        try:
            logger.info(f"Generation attempt {attempt}/{max_retries}")
            return generate_brd(data, output_path)
        except Exception as e:
            last_error = e
            logger.error(f"Attempt {attempt} failed: {e}")
            if attempt < max_retries:
                import time
                wait = attempt * 2
                logger.info(f"Retrying in {wait}s...")
                time.sleep(wait)

    raise RuntimeError(
        f"DOCX generation failed after {max_retries} attempts: {last_error}"
    )


# ==================== CLI Entry Point ====================

def main():
    """Точка входа для CLI-вызова: docx-gen.py <input_json> <output_docx>."""
    # Проверка зависимостей
    if len(sys.argv) == 2 and sys.argv[1] == '--check':
        try:
            from docx import Document
            print("✅ python-docx доступен")
            sys.exit(0)
        except ImportError:
            print("❌ python-docx не установлен. Выполните: pip install python-docx", file=sys.stderr)
            sys.exit(1)

    if len(sys.argv) < 3:
        print("Usage: docx-gen.py <input_json> <output_docx>", file=sys.stderr)
        print("       docx-gen.py --check   — проверка зависимостей", file=sys.stderr)
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]

    if not os.path.exists(input_path):
        print(f"❌ Файл не найден: {input_path}", file=sys.stderr)
        sys.exit(1)

    try:
        with open(input_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except json.JSONDecodeError as e:
        print(f"❌ Ошибка парсинга JSON: {e}", file=sys.stderr)
        sys.exit(1)
    except IOError as e:
        print(f"❌ Ошибка чтения файла: {e}", file=sys.stderr)
        sys.exit(1)

    try:
        result = generate_brd_with_retry(data, output_path)
        print(result)
    except Exception as e:
        print(f"❌ Ошибка генерации DOCX: {e}", file=sys.stderr)
        traceback.print_exc(file=sys.stderr)
        sys.exit(1)


if __name__ == '__main__':
    main()
