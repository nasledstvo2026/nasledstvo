"""Convert irina-npa-weekly.md to Word (.docx) with proper formatting."""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

md_path = "/home/user1/.openclaw/workspace/reports/irina-npa-weekly.md"
docx_path = "/home/user1/.openclaw/workspace/reports/irina-npa-weekly.docx"

with open(md_path, "r", encoding="utf-8") as f:
    lines = f.readlines()

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
font.color.rgb = RGBColor(0x22, 0x22, 0x22)

# Style heading fonts
for level, size, color in [(1, 18, 0x111111), (2, 15, 0x333333), (3, 13, 0x444444)]:
    hs = doc.styles[f'Heading {level}']
    hs.font.size = Pt(size)
    hs.font.color.rgb = RGBColor((color >> 16) & 0xFF, (color >> 8) & 0xFF, color & 0xFF)
    hs.font.bold = True

i = 0
in_table = False
table_rows = []

def add_table(doc, rows):
    if not rows:
        return
    ncols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            cell = table.cell(ri, ci)
            cell.text = cell_text.strip()
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    if ri == 0:
                        run.bold = True

while i < len(lines):
    line = lines[i].rstrip('\n')

    # Skip pure horizontal rules
    if re.match(r'^---+\s*$', line):
        i += 1
        continue

    # Table detection
    if '|' in line and re.match(r'^\s*\|', line):
        # Collect all table lines
        table_lines = []
        while i < len(lines) and '|' in lines[i] and re.match(r'^\s*\|', lines[i]):
            row_line = lines[i].strip()
            # Skip separator row
            if re.match(r'^\|[\s\-:|]+\|$', row_line):
                i += 1
                continue
            cells = [c.strip() for c in row_line.split('|')[1:-1]]
            table_lines.append(cells)
            i += 1
        add_table(doc, table_lines)
        doc.add_paragraph()  # spacing after table
        continue

    # Headings
    if line.startswith('### '):
        doc.add_heading(line[4:].strip(), level=3)
        i += 1
        continue
    if line.startswith('## '):
        doc.add_heading(line[3:].strip(), level=2)
        i += 1
        continue
    if line.startswith('# '):
        doc.add_heading(line[2:].strip(), level=1)
        i += 1
        continue

    # Empty line
    if not line.strip():
        i += 1
        continue

    # List items
    if re.match(r'^[-*]\s+', line):
        text = re.sub(r'^[-*]\s+', '', line)
        p = doc.add_paragraph(style='List Bullet')
        # Bold handling
        parts = re.split(r'\*\*(.*?)\*\*', text)
        for j, part in enumerate(parts):
            run = p.add_run(part)
            if j % 2 == 1:
                run.bold = True
        i += 1
        continue

    # Numbered list
    if re.match(r'^\d+\.\s+', line):
        text = re.sub(r'^\d+\.\s+', '', line)
        p = doc.add_paragraph(style='List Number')
        parts = re.split(r'\*\*(.*?)\*\*', text)
        for j, part in enumerate(parts):
            run = p.add_run(part)
            if j % 2 == 1:
                run.bold = True
        i += 1
        continue

    # Regular paragraph
    p = doc.add_paragraph()
    parts = re.split(r'\*\*(.*?)\*\*', line)
    for j, part in enumerate(parts):
        run = p.add_run(part)
        if j % 2 == 1:
            run.bold = True
    i += 1

doc.save(docx_path)
print(f"DOCX saved: {docx_path}")
