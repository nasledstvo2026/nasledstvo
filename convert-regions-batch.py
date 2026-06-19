"""Batch convert all region .md files to .docx"""
import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor

regions_dir = "/home/user1/.openclaw/workspace/reports/regions"

for fname in sorted(os.listdir(regions_dir)):
    if not fname.endswith(".md"):
        continue

    md_path = os.path.join(regions_dir, fname)
    docx_path = os.path.join(regions_dir, fname.replace(".md", ".docx"))

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    for level, size in [(1, 16), (2, 14), (3, 12)]:
        hs = doc.styles[f'Heading {level}']
        hs.font.size = Pt(size)
        hs.font.bold = True

    for line in lines:
        line = line.rstrip('\n')
        if re.match(r'^---+\s*$', line):
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif not line.strip():
            continue
        elif re.match(r'^[-*]\s+', line):
            text = re.sub(r'^[-*]\s+', '', line)
            p = doc.add_paragraph(style='List Bullet')
            parts = re.split(r'\*\*(.*?)\*\*', text)
            for j, part in enumerate(parts):
                run = p.add_run(part)
                if j % 2 == 1:
                    run.bold = True
        elif re.match(r'^\d+\.\s+', line):
            text = re.sub(r'^\d+\.\s+', '', line)
            p = doc.add_paragraph(style='List Number')
            parts = re.split(r'\*\*(.*?)\*\*', text)
            for j, part in enumerate(parts):
                run = p.add_run(part)
                if j % 2 == 1:
                    run.bold = True
        else:
            p = doc.add_paragraph()
            parts = re.split(r'\*\*(.*?)\*\*', line)
            for j, part in enumerate(parts):
                run = p.add_run(part)
                if j % 2 == 1:
                    run.bold = True

    doc.save(docx_path)
    print(f"OK: {fname} -> {fname.replace('.md', '.docx')}")

print("\nDone!")
